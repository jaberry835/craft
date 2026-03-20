"""
Template Engine for Policy Document Generation

Deterministic, server-side logic for filling Word (.docx) and
Excel (.xlsx) templates.  The LLM never touches Office files directly;
it passes structured field data and this module does the substitution.

Supported placeholder styles for .docx templates:
  1. Text placeholders:  ``{{field_name}}``  (typed directly in the document)
  2. Content Controls:   Plain-text or Rich-text controls whose **Tag**
     property is set to the field name.  Created via Word's Developer tab.

Both styles can coexist in the same template.
"""

import io
import logging
import re
from typing import Dict, Any, List
from xml.etree.ElementTree import Element

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from openpyxl import load_workbook

logger = logging.getLogger(__name__)

_PLACEHOLDER_RE = re.compile(r"\{\{(\w+)\}\}")

# Qualified names for structured document tags (content controls).
# qn() maps the "w:" prefix to the Word XML namespace internally —
# no network calls are made; namespaces are just identifier strings.
_SDT_TAG = qn("w:sdt")
_SDT_PR_TAG = qn("w:sdtPr")
_SDT_CONTENT_TAG = qn("w:sdtContent")
_TAG_TAG = qn("w:tag")
_ALIAS_TAG = qn("w:alias")
_VAL_ATTR = qn("w:val")


# ── Content Control (SDT) helpers ───────────────────────────────────

def _iter_sdts(element: Element):
    """Yield all <w:sdt> elements anywhere in the XML tree."""
    for sdt in element.iter(_SDT_TAG):
        yield sdt


def _sdt_field_name(sdt: Element) -> str | None:
    """Return the field name for a content control.

    Prefers the Tag property; falls back to Alias (Title).
    Returns None if neither is set.
    """
    pr = sdt.find(_SDT_PR_TAG)
    if pr is None:
        return None
    tag_el = pr.find(_TAG_TAG)
    if tag_el is not None:
        val = tag_el.get(_VAL_ATTR, "").strip()
        if val:
            return val
    alias_el = pr.find(_ALIAS_TAG)
    if alias_el is not None:
        val = alias_el.get(_VAL_ATTR, "").strip()
        if val:
            return val
    return None


def _set_sdt_text(sdt: Element, value: str):
    """Replace all text inside a content control with *value*.

    Keeps the first run's formatting and clears everything else.
    """
    content = sdt.find(_SDT_CONTENT_TAG)
    if content is None:
        return
    # Find all <w:r> (run) elements inside the content
    runs = list(content.iter(qn("w:r")))
    if not runs:
        return
    # Set text in first run's <w:t>, clear the rest
    first_t = runs[0].find(qn("w:t"))
    if first_t is None:
        # Create a <w:t> element if missing
        first_t = Element(qn("w:t"))
        runs[0].append(first_t)
    first_t.text = value
    first_t.set(qn("xml:space"), "preserve")
    for run in runs[1:]:
        t = run.find(qn("w:t"))
        if t is not None:
            t.text = ""


# ── Word (.docx) ────────────────────────────────────────────────────

def fill_docx(template_bytes: bytes, field_data: Dict[str, Any]) -> bytes:
    """Replace placeholders in a .docx template.

    Supports two placeholder styles:
      1. ``{{key}}`` text placeholders in paragraphs, tables,
         headers, and footers.
      2. Content Controls (structured document tags) whose Tag
         or Title matches a key in *field_data*.

    Returns the filled document as bytes.
    """
    doc = DocxDocument(io.BytesIO(template_bytes))

    def _replace_in_paragraph(paragraph):
        """Replace placeholders that may be split across multiple runs."""
        full_text = "".join(run.text for run in paragraph.runs)
        if not _PLACEHOLDER_RE.search(full_text):
            return
        new_text = _substitute(full_text, field_data)
        # Rewrite runs: put all text in the first run, clear the rest
        for i, run in enumerate(paragraph.runs):
            run.text = new_text if i == 0 else ""

    # Body paragraphs
    for para in doc.paragraphs:
        _replace_in_paragraph(para)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para)

    # Headers / footers
    for section in doc.sections:
        for header_footer in (section.header, section.footer):
            if header_footer is not None:
                for para in header_footer.paragraphs:
                    _replace_in_paragraph(para)

    # Content Controls (structured document tags)
    for sdt in _iter_sdts(doc.element.body):
        name = _sdt_field_name(sdt)
        if name and name in field_data:
            _set_sdt_text(sdt, str(field_data[name]))
    for section in doc.sections:
        for hf in (section.header, section.footer):
            if hf:
                for sdt in _iter_sdts(hf._element):
                    name = _sdt_field_name(sdt)
                    if name and name in field_data:
                        _set_sdt_text(sdt, str(field_data[name]))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Excel (.xlsx) ───────────────────────────────────────────────────

def fill_xlsx(template_bytes: bytes, field_data: Dict[str, Any]) -> bytes:
    """Replace ``{{key}}`` placeholders in an .xlsx template.

    Iterates all sheets and cells.  Returns the filled workbook as bytes.
    """
    wb = load_workbook(io.BytesIO(template_bytes))

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and _PLACEHOLDER_RE.search(cell.value):
                    cell.value = _substitute(cell.value, field_data)

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


# ── Shared helpers ──────────────────────────────────────────────────

def _substitute(text: str, field_data: Dict[str, Any]) -> str:
    """Replace all ``{{key}}`` occurrences with values from *field_data*."""
    def _repl(m):
        key = m.group(1)
        return str(field_data.get(key, m.group(0)))  # leave unreplaced if missing
    return _PLACEHOLDER_RE.sub(_repl, text)


def extract_placeholders(template_bytes: bytes, template_type: str) -> list[str]:
    """Return a sorted list of unique placeholder names found in a template."""
    if template_type == "docx":
        return _extract_docx_placeholders(template_bytes)
    if template_type == "xlsx":
        return _extract_xlsx_placeholders(template_bytes)
    return []


def _extract_docx_placeholders(data: bytes) -> list[str]:
    doc = DocxDocument(io.BytesIO(data))
    found: set[str] = set()

    # ── 1. Text-based {{field}} placeholders ──
    def _scan(paragraph):
        text = "".join(r.text for r in paragraph.runs)
        found.update(_PLACEHOLDER_RE.findall(text))

    for p in doc.paragraphs:
        _scan(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _scan(p)
    for s in doc.sections:
        for hf in (s.header, s.footer):
            if hf:
                for p in hf.paragraphs:
                    _scan(p)

    # ── 2. Content Controls (SDTs) ──
    for sdt in _iter_sdts(doc.element.body):
        name = _sdt_field_name(sdt)
        if name:
            found.add(name)
    for s in doc.sections:
        for hf in (s.header, s.footer):
            if hf:
                for sdt in _iter_sdts(hf._element):
                    name = _sdt_field_name(sdt)
                    if name:
                        found.add(name)

    return sorted(found)


def _extract_xlsx_placeholders(data: bytes) -> list[str]:
    wb = load_workbook(io.BytesIO(data))
    found: set[str] = set()
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    found.update(_PLACEHOLDER_RE.findall(cell.value))
    return sorted(found)
